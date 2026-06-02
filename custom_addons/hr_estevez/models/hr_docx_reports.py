# -*- coding: utf-8 -*-
"""
DOCX report generation for HR documents.
Business logic stays in hr_employee.py / hr_contract.py.
All DOCX rendering is isolated here.
"""
import io
import base64
import logging
from datetime import date, datetime
from PIL import Image, ImageDraw, ImageOps

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from odoo import models

_logger = logging.getLogger(__name__)

J = WD_ALIGN_PARAGRAPH.JUSTIFY
C = WD_ALIGN_PARAGRAPH.CENTER
R = WD_ALIGN_PARAGRAPH.RIGHT

# ---------------------------------------------------------------------------
# Utility builder — no Odoo dependencies
# ---------------------------------------------------------------------------

class _DocxBuilder:

    MESES_ES = {
        1: 'ENERO', 2: 'FEBRERO', 3: 'MARZO', 4: 'ABRIL',
        5: 'MAYO', 6: 'JUNIO', 7: 'JULIO', 8: 'AGOSTO',
        9: 'SEPTIEMBRE', 10: 'OCTUBRE', 11: 'NOVIEMBRE', 12: 'DICIEMBRE',
    }

    @classmethod
    def fmt_date(cls, d, sep=' DE '):
        if not d:
            return 'N/A'
        return f"{d.day:02d}{sep}{cls.MESES_ES[d.month]}{sep}{d.year}"

    @classmethod
    def fmt_today(cls):
        return cls.fmt_date(date.today())

    @classmethod
    def emp_address(cls, emp):
        parts = [
            emp.private_street or '',
            (' ' + emp.private_street2) if emp.private_street2 else '',
            (', ' + emp.private_colonia) if emp.private_colonia else '',
            (', ' + emp.private_city) if emp.private_city else '',
            (', ' + emp.private_state_id.name) if emp.private_state_id else '',
            (', C.P. ' + emp.private_zip) if emp.private_zip else '',
        ]
        return ''.join(parts).strip(', ') or 'N/A'

    @classmethod
    def company_address(cls, company):
        p = company.partner_id
        parts = [
            p.street or '',
            (', ' + p.street2) if p.street2 else '',
            (', ' + p.city) if p.city else '',
            (', ' + p.state_id.name) if p.state_id else '',
            (', C.P. ' + p.zip) if p.zip else '',
        ]
        return ''.join(parts).strip(', ') or 'N/A'

    @classmethod
    def entry_date_str(cls, employee):
        contracts = employee.contract_ids.filtered('date_of_entry')
        if not contracts:
            return 'N/A'
        return cls.fmt_date(contracts.sorted('date_of_entry')[0].date_of_entry)

    @staticmethod
    def remove_borders(table):
        tblPr = table._tbl.tblPr if table._tbl.tblPr is not None else OxmlElement('w:tblPr')
        tblBorders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            tblBorders.append(el)
        tblPr.append(tblBorders)

    @staticmethod
    def cell(cell, text, bold=False):
        para = cell.paragraphs[0]
        para.alignment = C
        run = para.add_run(str(text) if text else 'N/A')
        run.bold = bold

    @staticmethod
    def cell_left(cell, text, bold=False):
        run = cell.paragraphs[0].add_run(str(text) if text else 'N/A')
        run.bold = bold

    @staticmethod
    def p(doc, runs, align=J, space_before=0):
        """Add a paragraph with mixed bold/normal runs.
        runs: list of (text, bold) tuples.
        """
        para = doc.add_paragraph()
        para.alignment = align
        if space_before:
            para.paragraph_format.space_before = Pt(space_before)
        for text, bold in runs:
            r = para.add_run(str(text) if text else '')
            r.bold = bold
        return para

    @staticmethod
    def heading(doc, text, underline=False, space_before=8):
        para = doc.add_paragraph()
        para.alignment = C
        para.paragraph_format.space_before = Pt(space_before)
        run = para.add_run(text)
        run.bold = True
        if underline:
            run.underline = True
        return para

    @staticmethod
    def sig_table(doc, left_lines, right_lines, space_pt=56):
        """Two-column signature block."""
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(space_pt)
        tbl = doc.add_table(rows=1, cols=2)
        _DocxBuilder.remove_borders(tbl)
        for cell_obj, lines in ((tbl.rows[0].cells[0], left_lines),
                                (tbl.rows[0].cells[1], right_lines)):
            for i, (text, bold) in enumerate(lines):
                para = cell_obj.paragraphs[0] if i == 0 else cell_obj.add_paragraph()
                para.alignment = C
                run = para.add_run(str(text) if text else '')
                run.bold = bold
        return tbl

    @staticmethod
    def to_bytes(doc):
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def apply_default_typography(doc):
      """Apply a consistent base typography across all DOCX reports."""
      normal = doc.styles['Normal']
      normal.font.name = 'Times New Roman'
      normal.font.size = Pt(12)
      normal.paragraph_format.space_before = Pt(0)
      normal.paragraph_format.space_after = Pt(2)
      normal.paragraph_format.line_spacing = 1.15

    @staticmethod
    def circular_image_bytes(image_b64, size=320):
        """Return a circular PNG without zoom-cropping the original image."""
        img = Image.open(io.BytesIO(base64.b64decode(image_b64))).convert('RGBA')

        # Keep the full image visible by fitting inside the square canvas.
        fitted = ImageOps.contain(img, (size, size), method=Image.Resampling.LANCZOS)
        x = (size - fitted.width) // 2
        y = (size - fitted.height) // 2

        canvas = Image.new('RGBA', (size, size), (255, 255, 255, 0))
        canvas.paste(fitted, (x, y), fitted)

        mask = Image.new('L', (size, size), 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0, size - 1, size - 1), fill=255)

        out = Image.new('RGBA', (size, size), (255, 255, 255, 0))
        out.paste(canvas, (0, 0), mask)

        stream = io.BytesIO()
        out.save(stream, format='PNG')
        stream.seek(0)
        return stream

    @staticmethod
    def download(env, model, rec_id, filename, raw_bytes):
        att = env['ir.attachment'].create({
            'name': filename,
            'type': 'binary',
            'datas': base64.b64encode(raw_bytes).decode(),
            'res_model': model,
            'res_id': rec_id,
            'mimetype': (
                'application/vnd.openxmlformats-officedocument'
                '.wordprocessingml.document'
            ),
        })
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{att.id}?download=true',
            'target': 'new',
        }


# ---------------------------------------------------------------------------
# hr.employee DOCX reports
# ---------------------------------------------------------------------------

class HrEmployeeDocxReports(models.Model):
    _inherit = 'hr.employee'

    # --- Carta Patronal ---

    def action_download_carta_patronal_docx(self):
        self.ensure_one()
        b = _DocxBuilder
        doc = Document()
        b.apply_default_typography(doc)

        for sec in doc.sections:
            sec.top_margin = Inches(0.8)
            sec.bottom_margin = Inches(0.8)
            sec.left_margin = Inches(0.9)
            sec.right_margin = Inches(0.9)

        b.p(doc, [(f'TLALNEPANTLA DE BAZ A {b.fmt_today()}', True)], align=R)

        if self.image_1920:
            try:
                img_stream = b.circular_image_bytes(self.image_1920)
                para = doc.add_paragraph()
                para.alignment = R
                para.add_run().add_picture(img_stream, width=Inches(1.65))
            except Exception:
                pass

        doc.add_paragraph()
        b.p(doc, [('A QUIEN CORRESPONDA', True)])
        b.p(doc, [('PRESENTE:', True)])
        doc.add_paragraph()

        b.p(doc, [
            ('ME PERMITO INFORMAR A USTED, LOS DATOS QUE TENEMOS REGISTRADOS '
             'EN EL EXPEDIENTE DE EL C. ', False),
            (self.name.upper(), True),
            (' EL CUAL LABORA PARA LA EMPRESA Y CONTAMOS CON LOS SIGUIENTES '
             'DATOS EN NUESTROS REGISTROS:', False),
        ], align=J)
        doc.add_paragraph()

        t1 = doc.add_table(rows=5, cols=2)
        b.remove_borders(t1)
        for i, (label, val) in enumerate([
            ('FECHA DE INGRESO:', b.entry_date_str(self)),
            ('RFC:', self.rfc or 'N/A'),
            ('IMSS:', self.nss or 'N/A'),
            ('PUESTO:', self.job_id.name if self.job_id else 'N/A'),
            ('DOMICILIO PARTICULAR:', b.emp_address(self)),
        ]):
            b.cell_left(t1.rows[i].cells[0], label, bold=True)
            b.cell_left(t1.rows[i].cells[1], val)

        doc.add_paragraph()

        t2 = doc.add_table(rows=4, cols=2)
        b.remove_borders(t2)
        for i, (label, val) in enumerate([
            ('NOMBRE DE LA EMPRESA:', self.company_id.name or 'N/A'),
            ('DOMICILIO DE LA EMPRESA:', b.company_address(self.company_id)),
            ('RFC:', self.company_id.vat or 'No configurado'),
            ('REGISTRO PATRONAL:', 'C5350000 10 0'),
        ]):
            b.cell_left(t2.rows[i].cells[0], label, bold=True)
            b.cell_left(t2.rows[i].cells[1], val)

        b.p(doc, [('ATENTAMENTE', False)], align=C, space_before=42)
        b.p(doc, [('DEPARTAMENTO RECURSOS HUMANOS', False)], align=C)

        return b.download(self.env, 'hr.employee', self.id,
                          f'{self.name} - Carta Patronal.docx',
                          b.to_bytes(doc))

    # --- Convenio de Salida ---

    def action_download_convenio_salida_docx(self):
        self.ensure_one()
        b = _DocxBuilder
        emp = self.name.upper()
        co = self.company_id.name.upper()
        today = b.fmt_today()
        entry = b.entry_date_str(self)
        job = (self.job_id.name or 'N/A').upper()

        doc = Document()
        b.apply_default_typography(doc)
        # Convenio uses a different typography by business request.
        convenio_style = doc.styles['Normal']
        convenio_style.font.name = 'Calibri'
        convenio_style.font.size = Pt(10)
        for sec in doc.sections:
            sec.top_margin = Inches(1)
            sec.bottom_margin = Inches(1)
            sec.left_margin = Inches(1.2)
            sec.right_margin = Inches(1.2)

        b.p(doc, [
            ('CONVENIO QUE CELEBRAN POR UNA PARTE LA EMPRESA ', True),
            (co, True),
            (', REPRESENTADA POR EL LIC. EDWIN GONZALEZ SORIA EN SU CARÁCTER '
             'DE APODERADO LEGAL Y UNICO RESPONSABLE DE LA RELACIÓN LABORAL, '
             'EN LO SUCESIVO "LA EMPRESA" Y POR OTRA EL C. ', True),
            (emp, True),
            (', POR SU PROPIO DERECHO, EN SU CARÁCTER DE TRABAJADOR, EN LO '
             'SUCESIVO "EL TRABAJADOR", CONFORME A LOS SIGUIENTES, '
             'ANTECEDENTES, DECLARACIONES Y CLAUSULAS.', True),
        ], align=J)

        b.heading(doc, 'A N T E C E D E N T E S', underline=True)
        b.p(doc, [
            ('UNICO.', True),
            (f' - CON FECHA {entry} ', False),
            (co, True),
            (', Y EL C. ', False), (emp, True),
            (', CELEBRARON UN CONTRATO INDIVIDUAL DE TRABAJO CON EL OBJETO DE '
             'ESTABLECER Y REGULAR LA RELACIÓN LABORAL Y LAS CONDICIONES EN '
             'QUE SE IBA A PRESTAR ESTA.', False),
        ], align=J)

        b.heading(doc, 'D E C L A R A C I O N E S', underline=True)

        declaraciones = [
            ('I.', f'LA EMPRESA {co}, Y EL C. {emp}, MANIFIESTAN, QUE EN LA '
              'CELEBRACIÓN DEL PRESENTE CONVENIO, RECONOCEN MUTUAMENTE QUE NO '
              'HA EXISTIDO DOLO, MALA FE, ERROR, NI CUALQUIER OTRO VICIO DEL '
              'CONSENTIMIENTO QUE PUDIERA AFECTAR LA VALIDEZ DE ESTE '
              'INSTRUMENTO DE INEXISTENCIA O NULIDAD.'),
            ('II.', f'EL C. {emp}, EN SU CARACTER DE TRABAJADOR RECONOCE QUE '
              'EN EL CONTENIDO DEL PRESENTE CONVENIO, NO EXISTE NINGUNA '
              'RENUNCIA A SUS DE DERECHOS LABORALES NI MENOSCABO A ESTOS Y EL '
              'CONVENIO LO CELEBRA POR ASI CONVENIR A SUS INTERESES, POR SU '
              'PROPIO DERECHO, DE MANERA VOLUNTARIA, LIBRE Y ESPONTANEA SIN '
              'NINGUN TIPO DE PRESION, POR LO QUE NO EXISTE NINGUN VICIO DEL '
              'CONSENTIMIENTO QUE PUDIERA AFECTAR LA VALIDEZ DE ESTE CONVENIO '
              'DE INEXISTENCIA O NULIDAD.'),
            ('III.', f'LA EMPRESA {co}, Y EL C {emp}, MANIFIESTAN QUE EN SU '
              'CALIDAD DE TESTIGOS ESTAN PRESENTES EN EL DESARROLLO Y FIRMA '
              'DEL PRESENTE CONVENIO LA C. JESSICA CARINA CALLEJAS VAZQUEZ, '
              'LA C. CLAUDIA PEREZ LOPEZ.'),
            ('IV.', f'LA EMPRESA {co}, Y EL C {emp}, MANIFIESTAN QUE EL '
              'PRESENTE CONVENIO LO PACTAN DE ESTA FORMA POR RESULTAR MAS '
              'PRACTICO, YA QUE EN LA EJECUCION DEL MISMO SE REALIZA EN ESTE '
              'ACTO Y DE MANERA INMEDIATA, AHORRANDO EL TIEMPO QUE SE HUBIERA '
              'REQUERIDO PARA GENERAR ESTE TRAMITE ANTE EL CENTRO DE '
              'CONCILIACION LABORAL.'),
            ('V.', f'LA EMPRESA {co}, Y EL C {emp}, SE RECONOCEN '
              'RECÍPROCAMENTE SU PERSONALIDAD Y CAPACIDAD LEGAL PARA CELEBRAR '
              'EL PRESENTE CONVENIO POR LO TANTO SE OBLIGAN EN LOS TÉRMINOS '
              'DE LAS SIGUIENTES:'),
        ]
        for num, text in declaraciones:
            b.p(doc, [(num, True), (f' {text}', False)], align=J)

        b.heading(doc, 'C L Á U S U L A S', underline=True)

        clausulas = [
            ('PRIMERA.', f' DECLARAN, LA EMPRESA {co}, Y EL C {emp}, QUE EN '
              f'TERMINOS DE LA FRACCION I DEL ARTICULO 53 DE LA LEY FEDERAL '
              f'DEL TRABAJO, POR ASI CONVENIR A LOS INTERESES DEL C. {emp}, '
              f'POR MUTUO CONSENTIMIENTO, MANIFIESTAN Y RATIFICAN SU VOLUNTAD '
              f'DE DAR POR TERMINADA LA RELACION LABORAL QUE LOS UNIA '
              f'MEDIANTE EL CONTRATO INDIVIDUAL DE TRABAJO MENCIONADO EN EL '
              f'CAPITULO DE ANTECEDENTES DEL PRESENTE CONVENIO.'),
            ('SEGUNDA.', f' DECLARAN, EL C. {emp}, QUE POR ASI CONVENIR A '
              f'SUS INTERESES Y POR MUTUO CONSENTIMIENTO, ES SU LIBRE '
              f'VOLUNTAD DAR POR TERMINADA LA RELACIÓN LABORAL QUE LO HA '
              f'UNIDO CON LA EMPRESA {co}, HASTA EL DÍA {today} TERMINACIÓN '
              f'LABORAL QUE HACE CON FUNDAMENTO EN LO DISPUESTO POR LA '
              f'FRACCIÓN I DEL ARTÍCULO 53 DE LA LEY FEDERAL DEL TRABAJO.'),
            ('TERCERA.', f' DECLARAN, EL C. {emp}, QUE DURANTE EL TIEMPO QUE '
              f'LABORÓ PARA LA EMPRESA {co}, SIEMPRE RECIBIÓ DE ÉSTA DE '
              f'MANERA PUNTUAL, EL PAGO DE LOS SALARIOS ORDINARIOS Y '
              f'EXTRAORDINARIOS; SIEMPRE SE LE PAGARON LOS SÉPTIMOS DÍAS, '
              f'DÍAS FESTIVOS, Y HORAS EXTRAS EN LAS ESCASAS OCASIONES QUE '
              f'LOS LABORO, LOS SALARIOS DEVENGADOS, ASÍ COMO LAS VACACIONES, '
              f'PRIMA VACACIONAL Y AGUINALDO GENERADOS EN EL TIEMPO LABORADO, '
              f'COMO CONSECUENCIA NO SE LE ADEUDA NINGUNA SUMA POR DICHOS '
              f'CONCEPTOS NI POR NINGÚN OTRO.'),
            ('CUARTA.', f' DECLARAN: EL C. {emp}, QUE INICIO A PRESTAR SUS '
              f'SERVICIOS PERSONALES Y SUBORDINADOS, ÚNICAMENTE Y '
              f'EXCLUSIVAMENTE PARA LA EMPRESA {co}, A QUIEN RECONOCE COMO SU '
              f'ÚNICO PATRON, EN LOS TÉRMINOS Y CONDICIONES QUE A '
              f'CONTINUACIÓN SE MENCIONAN:'),
        ]
        for num, text in clausulas:
            b.p(doc, [(num, True), (text, False)], align=J)

        # Tabla fecha ingreso / categoría
        doc.add_paragraph()
        tbl = doc.add_table(rows=2, cols=2)
        b.cell_left(tbl.rows[0].cells[0], 'FECHA DE INGRESO', bold=True)
        b.cell_left(tbl.rows[0].cells[1], entry)
        b.cell_left(tbl.rows[1].cells[0], 'CATEGORÍA', bold=True)
        b.cell_left(tbl.rows[1].cells[1], job)
        doc.add_paragraph()

        clausulas2 = [
            ('QUINTA.', f' DECLARAN, LA EMPRESA {co}, POR CONDUCTO DEL LIC. '
              f'EDWIN GONZALEZ SORIA, QUIEN COMPARECE EN SU CARÁCTER DE '
              f'APODERADO LEGAL, QUE ESTÁ DE ACUERDO CON LO MANIFESTADO POR '
              f'EL C. {emp}, EN LAS CLÁUSULAS, SEGUNDA, TERCERA Y CUARTA Y '
              f'CONSECUENTE CON ELLO EXPRESA SU VOLUNTAD PARA DAR POR '
              f'TERMINADA POR MUTUO CONSENTIMIENTO LA RELACIÓN LABORAL CON '
              f'EL C. {emp}.'),
            ('SEXTA.', f' DECLARAN, LA EMPRESA {co}, Y EL C. {emp}, COMO '
              f'CONSECUENCIA DE LA TERMINACIÓN DEL CONTRATO INDIVIDUAL DE '
              f'TRABAJO, LA EMPRESA {co} POR CONDUCTO DEL LIC. EDWIN GONZALEZ '
              f'SORIA, ENTREGA AL C. {emp}, MEDIANTE TRANSFERENCIA QUE EN '
              f'COPIA SE ANEXA AL PRESENTE, LA CANTIDAD DE $0,000.000 '
              f'(CERO CERO CERO CERO PESOS 00/100 M.N.) COMO FINIQUITO DE '
              f'TODAS Y CADA UNA DE LAS PRESTACIONES A QUE TUVO DERECHO EL '
              f'C. {emp}, DURANTE LA RELACIÓN LABORAL QUE SE DA POR TERMINADA '
              f'EL DÍA {today}.'),
            ('SÉPTIMA.', f' DECLARAN: EL C. {emp}, QUE ESTA CONFORME CON LA '
              f'CANTIDAD Y CONCEPTOS CITADOS POR LA EMPRESA ESTEVEZ.JOR '
              f'SERVICIOS, S.A. DE C.V., EN LA CLÁUSULA ANTERIOR YA QUE EN '
              f'SU OPORTUNIDAD SIEMPRE RECIBIÓ DE MANERA ININTERRUMPIDA EL '
              f'PAGO DE LOS SÉPTIMOS DÍAS, DÍAS FESTIVOS, Y HORAS EXTRAS, '
              f'LOS SALARIOS DEVENGADOS, ASÍ COMO LAS VACACIONES, PRIMA '
              f'VACACIONAL Y AGUINALDO GENERADOS EN EL TIEMPO LABORADO.'),
            ('OCTAVA.', f' DECLARA: EL C. {emp}, QUE EL PRESENTE DOCUMENTO NO '
              f'REPRESENTA NINGUNA RENUNCIA A SUS DERECHOS O MENOSCABO A '
              f'ESTOS, POR QUE LOS DERECHOS QUE GENERO LE ESTAN SIENDO '
              f'PAGADOS CONFORME A LA LEY FEDERAL DEL TRABAJO, MOTIVO POR EL '
              f'CUAL NO SE RESERVA NINGUNA ACCIÓN O DERECHO QUE EJERCITAR EN '
              f'CONTRA DE LA EMPRESA {co}, A QUIEN RECONOCE COMO SU UNICO '
              f'PATRON, POR NINGÚN CONCEPTO DE TIPO LABORAL, CIVIL, PENAL, '
              f'ADMINISTRATIVO O CUALQUIER OTRO.'),
            ('NOVENA.', f' DECLARAN, LA EMPRESA {co} Y EL C. {emp}, PARA '
              f'EFECTOS DE CUMPLIR CON LOS REQUISITOS DE VALIDEZ QUE REFIERE '
              f'EL SEGUNDO PARRAFO DEL ARTICULO 33 DE LA LEY FEDERAL DEL '
              f'TRABAJO, QUE LA RELACION CIRCUNSTANCIADA DE LOS HECHOS QUE '
              f'MOTIVAN EL PRESENTE CONVENIO ASI COMO LOS DERECHOS '
              f'COMPRENDIDOS EN ÉL, SE ENCUENTRAN EXPRESADOS EN EL CAPITULO '
              f'DE ANTECEDENTES, DECLARACIONES Y EN TODAS Y CADA UNA DE LAS '
              f'CLAUSULAS QUE SE MENCIONAN EN EL CUERPO DEL PRESENTE CONVENIO.'),
            ('DÉCIMA.', f' SEÑALAN LAS PARTES COMO DOMICILIOS PARA EFECTO DE '
              f'OIR Y RECIBIR NOTIFICACIONES LA EMPRESA {co} EL UBICADO EN '
              f'FILIBERTO GOMEZ N°46, CENTRO INDUSTRIAL TLALNEPANTLA, '
              f'TLALNEPANTLA DE BAZ, ESTADO DE MÉXICO. C.P. 54030. Y EL C. '
              f'{emp} EL UBICADO EN REVOLUCIÓN #296, LA ROMANA TLALNEPANTLA '
              f'DE BAZ, MÉXICO C.P. 54030.'),
            ('DÉCIMA PRIMERA.', f' DECLARAN, LA EMPRESA {co} Y EL C. {emp}, '
              f'QUE EL PRESENTE CONVENIO POR NO CONTENER CLÁUSULAS CONTRARIAS '
              f'A LA MORAL O AL DERECHO Y POR ESTAR DENUNCIADO CONFORME A LO '
              f'PREVISTO POR LOS ARTÍCULOS 33, Y 53, FRACCIÓN I, DE LA LEY '
              f'FEDERAL DEL TRABAJO VIGENTE, SE OBLIGAN A ESTAR Y PASAR POR '
              f'EL EN TODO TIEMPO Y LUGAR COMO SI SE TRATARA DE COSA JUZGADA.'),
        ]
        for num, text in clausulas2:
            b.p(doc, [(num, True), (text, False)], align=J)

        b.p(doc, [
            (f'LEÍDO QUE FUE EN SU TOTALIDAD EL CONTENIDO DEL PRESENTE '
             f'DOCUMENTO POR LA EMPRESA {co} EN SU CALIDAD DE PATRON, EL C. '
             f'{emp} EN SU CALIDAD DE TRABAJADOR, Y LAS C. JESSICA CARINA '
             f'CALLEJAS VAZQUEZ Y CLAUDIA PEREZ LOPEZ EN SU CALIDAD DE '
             f'TESTIGOS, LO FIRMAN Y RATIFICAN EN EL MUNICIPIO DE '
             f'TLALNEPANTLA DE BAZ, ESTADO DE MEXICO A {today}.', False),
        ], align=J)

        b.sig_table(doc,
            left_lines=[
                (co, True), ('RESPONSABLE DE LA RELACIÓN DE TRABAJO', False),
                ('', False),
                ('LIC. EDWIN GONZALEZ SORIA', True), ('APODERADO LEGAL.', True),
                ('', False),
                ('TESTIGO', False), ('', False),
                ('JESSICA CARINA CALLEJAS VAZQUEZ', False),
            ],
            right_lines=[
                ('EL TRABAJADOR', True), ('.', False),
                ('', False),
                (emp, True), ('.', False),
                ('', False),
                ('TESTIGO', False), ('', False),
                ('CLAUDIA PEREZ LOPEZ', False),
            ],
        )

        return b.download(self.env, 'hr.employee', self.id,
                          f'Convenio de Salida - {self.name}.docx',
                          b.to_bytes(doc))

    # --- Oficio de Remisión ---

    def action_download_remision_docx(self):
        self.ensure_one()
        b = _DocxBuilder
        emp = self.name.upper()
        co = self.company_id.name.upper()
        today = f"{date.today().day} DE {b.MESES_ES[date.today().month]} DE {date.today().year}"

        doc = Document()
        b.apply_default_typography(doc)
        remision_style = doc.styles['Normal']
        remision_style.font.name = 'Arial'
        remision_style.font.size = Pt(10)
        remision_style.paragraph_format.space_before = Pt(0)
        remision_style.paragraph_format.space_after = Pt(0)
        remision_style.paragraph_format.line_spacing = 1.0
        for sec in doc.sections:
            sec.top_margin = Inches(0.55)
            sec.bottom_margin = Inches(0.55)
            sec.left_margin = Inches(0.9)
            sec.right_margin = Inches(0.9)

        # Header + addressed block to emulate the visual distribution in the sample.
        tbl = doc.add_table(rows=2, cols=2)
        b.remove_borders(tbl)
        tbl.autofit = False
        tbl.columns[0].width = Inches(3.05)
        tbl.columns[1].width = Inches(2.95)

        right = tbl.rows[0].cells[1]
        for i, (label, val, bold_val) in enumerate([
            ('Empresa:', co, True),
            ('Trabajador:', emp, True),
            ('Actividad De La Empresa:', 'Instalación de Fibra Óptica', False),
            ('Domicilio:', b.company_address(self.company_id).upper(), False),
        ]):
            para = right.paragraphs[0] if i == 0 else right.add_paragraph()
            para.alignment = J
            para.paragraph_format.line_spacing = 1.0
            r1 = para.add_run(f'{label} ')
            r1.bold = True
            r2 = para.add_run(val)
            r2.bold = bold_val

        left = tbl.rows[1].cells[0]
        for i, (text, bold) in enumerate([
            ('H. PRIMER TRIBUNAL LABORAL', True),
            ('DE LA REGION JUDICIAL', True),
            ('TLALNEPANTLA.', True),
            ('', False),
            ('H. CENTRO DE CONCILIACION LABORAL', True),
            ('CON SEDE EN TLALNEPANTLA ESTADO', True),
            ('DE MEXICO.', True),
            ('P R E S E N T E.', True),
        ]):
            para = left.paragraphs[0] if i == 0 else left.add_paragraph()
            para.alignment = J
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(text)
            run.bold = bold

        doc.add_paragraph()
        doc.add_paragraph()

        p1 = b.p(doc, [
            ('LIC. EDWIN GONZALEZ SORIA,', True),
            (' en mi carácter de Apoderado Legal ', False),
            (co, True),
            (', y el C. ', False),
            (emp, True),
            (', en mi carácter de trabajador, AMBOS señalando como domicilio '
             'para oír y recibir notificaciones el ubicado en FILIBERTO GOMEZ '
             'NÚMERO 46 INTERIOR 101 (JURÍDICO), COLONIA CENTRO INDUSTRIAL '
             'TLALNEPANTLA, MUNICIPIO DE TLALNEPANTLA DE BAZ, ESTADO DE MÉXICO, '
             'CÓDIGO POSTAL 54030, ante usted respetuosamente comparecemos a '
             'exponer:', False),
        ], align=J)
        p1.paragraph_format.line_spacing = 1.0
        p1.paragraph_format.space_after = Pt(18)

        p2 = b.p(doc, [
            ('Que, por medio del presente ocurso, venimos a presentar el '
             'convenio que hemos celebrado por medio del cual y de común '
             'acuerdo damos por terminada la relación y/o vínculo laboral que '
             'unía en términos de la fracción I del artículo 53 de la ley '
             'federal del trabajo, a la ', False),
            ('empresa ', True), (co, True),
            (', con el C. ', False), (emp, False),
            (', en su carácter de trabajador.', True),
        ], align=J)
        p2.paragraph_format.line_spacing = 1.0
        p2.paragraph_format.space_after = Pt(18)

        p3 = b.p(doc, [
            ('Ambas partes nos comprometemos a ratificar el convenio que se '
             'anexa el día y hora que esta autoridad nos señale fecha para tal '
             'efecto.', False),
        ], align=J)
        p3.paragraph_format.line_spacing = 1.0
        p3.paragraph_format.space_after = Pt(20)

        p4 = b.p(doc, [
            ('Por lo anteriormente expuesto, a este H. Tribunal, atentamente '
             'solicito.', False),
        ], align=C)
        p4.paragraph_format.line_spacing = 1.0

        p5 = b.p(doc, [
            ('Único. -', True),
            (' Tenernos por presentado en términos del presente escrito.',
             False),
        ], align=C)
        p5.paragraph_format.line_spacing = 1.0
        p5.paragraph_format.space_after = Pt(12)

        p6 = b.p(doc, [
            (f'Tlalnepantla de Baz Estado De México, a ', False),
            (today, True),
        ], align=C)
        p6.paragraph_format.line_spacing = 1.0

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(34)
        sig_tbl = doc.add_table(rows=1, cols=2)
        b.remove_borders(sig_tbl)
        sig_tbl.autofit = False
        sig_tbl.columns[0].width = Inches(3.05)
        sig_tbl.columns[1].width = Inches(2.95)

        left_sig = sig_tbl.rows[0].cells[0]
        for i, (text, bold) in enumerate([
            ('LIC. EDWIN GONZALEZ SORIA', True),
            ('APODERADO LEGAL', True),
            ('ESTEVEZ.JOR SERVICIOS, S.A. DE C.V.', True),
        ]):
            para = left_sig.paragraphs[0] if i == 0 else left_sig.add_paragraph()
            para.alignment = C
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(text)
            run.bold = bold

        right_sig = sig_tbl.rows[0].cells[1]
        for i, (text, bold) in enumerate([
            (f'EL C {emp}.', True),
            ('Trabajador', False),
        ]):
            para = right_sig.paragraphs[0] if i == 0 else right_sig.add_paragraph()
            para.alignment = C
            para.paragraph_format.line_spacing = 1.0
            run = para.add_run(text)
            run.bold = bold

        return b.download(self.env, 'hr.employee', self.id,
                          f'{self.name} - Oficio de Remisión.docx',
                          b.to_bytes(doc))


# ---------------------------------------------------------------------------
# hr.contract DOCX reports
# ---------------------------------------------------------------------------

class HrContractDocxReports(models.Model):
    _inherit = 'hr.contract'

    # --- Contrato Individual de Trabajo ---

    def action_download_contract_docx(self):
        self.ensure_one()
        b = _DocxBuilder
        emp = self.employee_id
        emp_name = emp.name.upper() if emp else 'N/A'
        co = self.company_id.name.upper()
        job = (emp.job_id.name.upper() if emp.job_id else 'N/A')
        date_str = b.fmt_date(self.date_start) if self.date_start else 'N/A'

        gender_map = {
            'male': 'MASCULINO', 'female': 'FEMENINO', 'indistinct': 'INDISTINTO',
        }
        marital_map = {
            'single': 'SOLTERO', 'married': 'CASADO', 'cohabitant': 'UNIÓN LIBRE',
            'widower': 'VIUDO', 'divorced': 'DIVORCIADO',
        }
        gender = gender_map.get(emp.gender, 'N/A') if emp else 'N/A'
        marital = marital_map.get(emp.marital, 'N/A') if emp else 'N/A'
        age = str(emp.age) if emp and emp.age else 'N/A'
        nationality = emp.get_nationality() if emp else 'N/A'
        curp = emp.curp or 'N/A'
        rfc = emp.rfc or 'N/A'
        address = b.emp_address(emp) if emp else 'N/A'

        doc = Document()
        b.apply_default_typography(doc)
        for sec in doc.sections:
            sec.top_margin = Inches(0.5)
            sec.bottom_margin = Inches(0.5)
            sec.left_margin = Inches(0.7)
            sec.right_margin = Inches(0.7)

        b.heading(doc, 'CONTRATO INDIVIDUAL DE TRABAJO')
        b.p(doc, [
            ('CONTRATO INDIVIDUAL DE TRABAJO QUE CONFORME A LO ESTABLECIDO '
             'POR LOS ARTÍCULOS 24 Y 25 DE LA LEY FEDERAL DEL TRABAJO, '
             'CELEBRAN POR UNA PARTE LA EMPRESA ', False),
            (co, True),
            (', REPRESENTADA LEGALMENTE POR EL LIC. ', False),
            ('EDWIN GONZALEZ SORIA', True),
            (', EN SU CARÁCTER DE APODERADO LEGAL A QUIÉN EN LO SUCESIVO Y '
             'PARA EFECTOS DE ESTE CONTRATO SE LE DENOMINARÁ ', False),
            ('"EL PATRON"', True),
            (', Y POR LA OTRA POR SU PROPIO DERECHO, EL/LA C. ', False),
            (emp_name, True),
            (', A QUIÉN EN LO SUCESIVO Y PARA EFECTOS DE ESTE CONTRATO SE LE '
             'DENOMINARÁ ', False),
            ('"EL TRABAJADOR"', True),
            (', AL TENOR DE LAS SIGUIENTES DECLARACIONES Y CLÁUSULAS:', False),
        ], align=J)

        b.heading(doc, 'DECLARACIONES:')
        b.p(doc, [('A. ', False), ('DE "EL PATRON":', True)])

        patron_decl = [
            ('1.', ' QUE ES UNA SOCIEDAD ANÓNIMA DE CAPITAL VARIABLE '
              'DEBIDAMENTE CONSTITUIDA Y LEGALMENTE EXISTENTE DE ACUERDO CON '
              'LAS LEYES MEXICANAS, TAL Y COMO SE ACREDITA EN TÉRMINOS DE '
              'INSTRUMENTO NOTARIAL NÚMERO 46,404, DE FECHA 4 DE OCTUBRE DE '
              '2011 PASADO ANTE LA FE DEL LIC. JOSE ORLANDO PADILLA BENITEZ, '
              'NOTARIO PÚBLICO NO. 30 DEL ESTADO DE MÉXICO.'),
            ('2.', ' SU REPRESENTANTE LEGAL CUENTA CON PLENAS Y SUFICIENTES '
              'FACULTADES PARA CELEBRAR EL PRESENTE CONTRATO, TAL Y COMO LO '
              'ACREDITA CON EL INSTRUMENTO NOTARIAL NÚMERO 18,435, DE FECHA '
              '07 DE DICIEMBRE DE 2020, PASADO ANTE LA FE DEL LICENCIADO JUAN '
              'CARLOS FRANCISCO DIAZ PONCE DE LEON, NOTARIO PÚBLICO NÚMERO '
              '209 DE LA CIUDAD DE MÉXICO.'),
            ('3.', ' QUE TIENE SU DOMICILIO EN: CALLE FILIBERTO GOMEZ NUMERO '
              '46 COLONIA CENTRO INDUSTRIAL TLALNEPANTLA, C.P. 54030, '
              'MUNICIPIO DE TLALNEPANTLA, ESTADO DE MEXICO.'),
            ('4.', ' QUE SU OBJETO SOCIAL CONSISTE EN LA INSTALACION Y '
              'MANTEMIENTO DE FIBRA ÓPTICA ENTRE OTROS.'),
            ('5.', ' QUE QUIEN SUSCRIBE POR "EL PATRON", CUENTA CON '
              'BASTANTES, AMPLIAS Y SUFICIENTES FACULTADES PARA LA '
              'CELEBRACION DE ESTE CONTRATO, LAS CUALES AL DIA DE HOY NO LE '
              'HAN SIDO REVOCADAS, RESTRINGIDAS, SUSPENDIDAS O DE ALGUNA OTRA '
              'MANERA LIMITADAS.'),
            ('6.', ' QUE CUENTA CON EL PERSONAL NECESARIO PARA EL DESARROLLO '
              'DE SUS OPERACIONES NORMALES, Y REQUIERE CONTRATAR PERSONAL QUE '
              'CUENTE CON LA DISPOSICIÓN PARA DESEMPEÑAR LOS TRABAJOS '
              'RELACIONADOS CON EL OBJETO SOCIAL DE LA EMPRESA.'),
        ]
        for num, txt in patron_decl:
            b.p(doc, [(num, True), (txt, False)], align=J)

        b.p(doc, [('B. ', False), ('DE "EL TRABAJADOR":', True)], space_before=6)
        b.p(doc, [
            ('1.', True),
            (f' QUE CUENTA CON LA DISPOSICIÓN NECESARIA PARA DESEMPEÑAR LOS '
             f'SERVICIOS EN EL PUESTO REQUERIDO, DE: ', False),
            (job, True),
        ], align=J)
        b.p(doc, [('2.', True), (' QUE SUS GENERALES SON LOS SIGUIENTES:', False)], align=J)

        # Generales table
        doc.add_paragraph()
        gen_tbl = doc.add_table(rows=7, cols=2)
        for i, (lbl, val) in enumerate([
            ('NACIONALIDAD', nationality),
            ('EDAD', f'{age} Años'),
            ('SEXO', gender),
            ('ESTADO CIVIL', marital),
            ('CURP', curp),
            ('RFC', rfc),
            ('DOMICILIO', address),
        ]):
            b.cell_left(gen_tbl.rows[i].cells[0], lbl, bold=True)
            b.cell_left(gen_tbl.rows[i].cells[1], val)
        doc.add_paragraph()

        b.p(doc, [('3.', True), (' QUE BAJO PROTESTA DE DECIR VERDAD, DECLARA '
            'QUE TODA LA INFORMACIÓN QUE HA PROPORCIONADO DE MANERA ORAL O '
            'ESCRITA A "EL PATRÓN" ES CORRECTA Y FIDEDIGNA.', False)], align=J)
        b.p(doc, [('4.', True), (' QUE LE HAN ENTERADO Y EXPLICADO LAS '
            'CONDICIONES DE TRABAJO Y COMO CONSECUENCIA ES SU EXPRESA VOLUNTAD '
            'CELEBRAR EL PRESENTE CONTRATO EN LOS TÉRMINOS Y CONDICIONES '
            'DESCRITOS EN EL CUERPO DEL PRESENTE INSTRUMENTO.', False)], align=J)

        b.p(doc, [('EXPUESTAS LAS ANTERIORES DECLARACIONES, AMBAS PARTES SE '
            'SOMETEN A LAS SIGUIENTES:', False)], align=J)
        b.p(doc, [('DECLARAN "LAS PARTES":', True)])
        b.p(doc, [('ÚNICO.', True), (' QUE AMBAS PARTES SE RECONOCEN EL '
            'CARÁCTER CON EL QUE COMPARECEN Y MANIFIESTAN QUE ES SU DESEO '
            'CELEBRAR EL PRESENTE CONTRATO EN LOS TÉRMINOS QUE A CONTINUACIÓN '
            'SE ESTIPULAN.', False)], align=J)

        b.heading(doc, 'CLÁUSULAS:')

        clausulas = [
            ('PRIMERA.', ' PARA EFECTOS DE COMPRENSIÓN DEL PRESENTE CONTRATO, '
              'SE DENOMINARÁ EN LO SUCESIVO A LA LEY FEDERAL DEL TRABAJO COMO '
              '"LA LEY", AL REFERIRSE AL PRESENTE DOCUMENTO COMO "EL CONTRATO", '
              'AL REGLAMENTO INTERIOR DE TRABAJO "REGLAMENTO" Y DE MANERA '
              'CONJUNTA A LOS QUE SUSCRIBEN COMO "LAS PARTES".'),
            ('SEGUNDA.', ' EL PRESENTE CONTRATO REGIRÁ EN TODO TIEMPO Y LUGAR '
              'LO NO ESTIPULADO EN EL MISMO POR EL "REGLAMENTO" Y POR LO '
              'ESTABLECIDO EN "LA LEY".'),
            ('TERCERA.', f' "EL PATRÓN", EN ESTE ACTO CONTRATA A "EL '
              f'TRABAJADOR" BAJO LA MODALIDAD DE CAPACITACIÓN INICIAL PARA QUE '
              f'DESEMPEÑE EL PUESTO DE {job} ASÍ COMO EN CUALQUIER OTRA '
              f'ACTIVIDAD ANEXA, CONEXA O RELACIONADA QUE LE ENCOMIENDE EL '
              f'PATRÓN, COMPATIBLE CON EL FIN DE SU CATEGORÍA, CAPACITACIÓN Y '
              f'DE SUS APTITUDES.'),
            ('CUARTA.', ' "EL TRABAJADOR" RECONOCE Y ESTA CONFORME, DADA LA '
              'MODALIDAD DE CONTRATACIÓN DE CAPACITACIÓN INICIAL, QUE A PARTIR '
              'DE LA SUSCRIPCIÓN DEL PRESENTE CONTRATO, LA RELACIÓN LABORAL '
              'TENDRÁ UNA VIGENCIA POR TIEMPO DETERMINADO DE CERO DIAS '
              'NATURALES, CONTADOS A PARTIR DE LA FECHA DE SU FIRMA.'),
            ('QUINTA.', f' "EL TRABAJADOR" SE OBLIGA A PRESTAR A "EL PATRÓN", '
              f'BAJO SU DIRECCIÓN Y DEPENDENCIA, SUS SERVICIOS PERSONALES '
              f'COMO: {job}.'),
            ('SEXTA.', ' "EL TRABAJADOR" DEBE DESEMPEÑAR DICHAS ACTIVIDADES EN '
              'EL DOMICILIO UBICADO EN CALLE FILIBERTO GÓMEZ NUMERO 46 COLONIA '
              'CENTRO INDUSTRIAL TLALNEPANTLA, C.P. 54030, MUNICIPIO DE '
              'TLALNEPANTLA, ESTADO DE MEXICO., Y/O INDISTINTAMENTE EN '
              'CUALQUIER DOMICILIO O LUGAR QUE SE LE INDIQUE POR "EL PATRÓN".'),
            ('SÉPTIMA.', ' "EL TRABAJADOR" DESARROLLARA SUS FUNCIONES EN '
              'CUALQUIER LUGAR QUE LE SEA ASIGNADO POR "EL PATRÓN", DE ACUERDO '
              'A LAS NECESIDADES DE LA EMPRESA.'),
            ('OCTAVA.', ' LA DURACIÓN DE LA JORNADA DE TRABAJO SERÁ DE 8 HORAS '
              'DIARIAS EN UN HORARIO QUE SE COMPRENDERÁ DE LAS 08:00 AM A LAS '
              '17:00 PM O BIEN DE LAS 9:00 AM A 18:00 HORAS O DE 7:00 AM A '
              '16:00 HORAS, DE LUNES A VIERNES Y SÁBADOS DE 8:00 AM A 14:00 '
              'HORAS. DURANTE LA JORNADA CONTINUA SE CONCEDERÁ AL TRABAJADOR '
              'UN DESCANSO DE UNA HORA DE 14:00 PM A 15:00 PM.'),
            ('NOVENA.', ' EL SALARIO MENSUAL QUE RECIBIRÁ "EL TRABAJADOR" POR '
              'SUS SERVICIOS SERÁ DE $0,000.000 (CERO CERO CERO CERO PESOS '
              '00/100 M.N.), MENOS LOS DESCUENTOS DE LEY. EL PAGO SE HARÁ EN '
              'FORMA QUINCENAL LOS DÍAS QUINCE Y ÚLTIMO DÍA DE CADA MES.'),
            ('DECIMA.', ' PREVISIÓN SOCIAL "PPP". ADEMÁS DEL SALARIO, EL '
              'EMPLEADO PERCIBIRÁ LAS PRESTACIONES DE PREVISIÓN SOCIAL DE '
              'ACUERDO AL PLAN PRIVADO DE PENSIONES QUE TIENE CONTRATADO LA '
              'EMPRESA, DEBIDAMENTE REGISTRADO ANTE LA CONSAR.'),
            ('DECIMA PRIMERA.', ' "EL TRABAJADOR" TENDRÁ DERECHO POR CADA SEIS '
              'DÍAS DE LABORES A DESCANSAR UNO CON EL PAGO DE SALARIO DIARIO '
              'CORRESPONDIENTE. DE ACUERDO AL ARTÍCULO 74 DE LA LEY FEDERAL '
              'DEL TRABAJO, SERÁN DÍAS DE DESCANSO OBLIGATORIO: EL 1 DE ENERO, '
              '1 DE MAYO, 16 DE SEPTIEMBRE, 25 DE DICIEMBRE, ENTRE OTROS.'),
            ('DECIMA SEGUNDA.', ' CUANDO "EL TRABAJADOR" POR RAZONES '
              'ADMINISTRATIVAS TENGA QUE LABORAR EL DÍA DOMINGO, "EL PATRÓN" '
              'LE PAGARÁ, ADEMÁS DE SU SALARIO, 25% COMO PRIMA DOMINICAL SOBRE '
              'EL SALARIO ORDINARIO DEVENGADO.'),
            ('DÉCIMA TERCERA.', ' "EL TRABAJADOR" TENDRÁ DERECHO A DISFRUTAR '
              'DE UN PERIODO ANUAL DE VACACIONES SEGÚN LO ESTABLECIDO EN EL '
              'ARTÍCULO 76 DE "LA LEY", ASÍ COMO A LA PRIMA VACACIONAL '
              'EQUIVALENTE AL 25% DEL IMPORTE PAGADO POR CONCEPTO DE '
              'VACACIONES.'),
            ('DÉCIMA CUARTA.', ' "EL TRABAJADOR" TENDRÁ DERECHO A RECIBIR '
              'ANTES DEL DÍA 20 DE DICIEMBRE DE CADA AÑO EL IMPORTE '
              'CORRESPONDIENTE A QUINCE DÍAS DE SALARIO COMO PAGO DEL '
              'AGUINALDO.'),
            ('DÉCIMA QUINTA.', ' "EL TRABAJADOR" DEBERÁ PRESENTARSE '
              'PUNTUALMENTE A SUS LABORES EN EL HORARIO DE TRABAJO ESTABLECIDO '
              'Y FIRMAR LAS LISTAS DE ASISTENCIA O CHECAR SU TARJETA EN EL '
              'RELOJ CHECADOR DIARIAMENTE.'),
            ('DECIMA SEXTA.', ' SIN PERJUICIO DE LO ESTIPULADO EN EL PRESENTE '
              'CONTRATO, CUANDO "EL TRABAJADOR" TENGA MAS DE TRES FALTAS '
              'INJUSTIFICADAS EN UN PERIODO DE TREINTA DIAS SE PROCEDERA A '
              'APLICAR EL SUPUESTO CONTEMPLADO EN EL ARTICULO 47, FRACCION X '
              'DE LA LEY.'),
            ('DÉCIMA SEPTIMA.', ' "EL PATRON" SE OBLIGA A CAPACITAR O A '
              'ADIESTRAR A "EL TRABAJADOR" DE ACUERDO A LOS PLANES Y PROGRAMAS '
              'QUE EXISTEN O SE ESTABLEZCAN. "EL TRABAJADOR" SE OBLIGA A '
              'OBSERVAR LAS DISPOSICIONES EN MATERIA DE SEGURIDAD, SALUD Y '
              'MEDIO AMBIENTE DE TRABAJO.'),
            ('DÉCIMA OCTAVA.', ' "EL TRABAJADOR" ACEPTA SOMETERSE A LOS '
              'EXÁMENES MÉDICOS QUE PERIÓDICAMENTE ESTABLEZCA "EL PATRÓN" EN '
              'LOS TÉRMINOS DEL ARTICULO 134 FRACCIÓN X DE "LA LEY".'),
            ('DÉCIMA NOVENA.', ' "EL PATRON" Y "EL TRABAJADOR" SE OBLIGAN EN '
              'TODO TIEMPO Y LUGAR A CUMPLIR CON LAS OBLIGACIONES QUE "LA LEY" '
              'LES IMPONEN EN TÉRMINOS DE LO ESTABLECIDO POR LOS ARTÍCULOS '
              '132 AL 135.'),
            ('VIGÉSIMA.', ' "EL TRABAJADOR" SABE Y ENTIENDE QUE TODA LA '
              'INFORMACIÓN QUE EXISTE Y LLEGUE A EXISTIR DE "EL PATRON" TIENE '
              'EL CARÁCTER DE "CONFIDENCIAL". "EL TRABAJADOR" SE OBLIGA A NO '
              'EXTRAER, HACER USO, COMUNICAR, DIVULGAR NI REVELAR POR NINGÚN '
              'MEDIO TAL INFORMACIÓN.'),
            ('VIGESIMA PRIMERA.', ' "EL TRABAJADOR" SE OBLIGA DURANTE LA '
              'VIGENCIA DE ESTE CONTRATO, ASÍ COMO DESPUÉS DE SU TERMINACIÓN '
              'O RESCISIÓN, A NO DIVULGAR NI UTILIZAR CUALQUIER ASPECTO O '
              'INFORMACIÓN RELACIONADA CON LOS NEGOCIOS, ACTIVIDADES Y '
              'OPERACIONES DE "EL PATRON" O DE LOS CLIENTES DE ÉSTE.'),
            ('VIGESIMA SEGUNDA.', ' EN CASO DE QUE "EL TRABAJADOR" INCUMPLA '
              'CON LAS OBLIGACIONES A SU CARGO, PREVISTAS EN LAS CLAUSULAS QUE '
              'ANTECEDEN; PAGARÁ A "EL PATRON" UNA INDEMNIZACIÓN EQUIVALENTE '
              'AL MONTO DE LOS DAÑOS Y PERJUICIOS QUE POR ESTE CONCEPTO SE '
              'GENERASEN.'),
            ('VIGESIMA TERCERA.', ' "EL TRABAJADOR", DURANTE EL DESEMPEÑO DE '
              'SUS FUNCIONES SE RESPONSABILIZA POR CUALQUIER ACTO U OMISIÓN, '
              'INTENCIONAL O NEGLIGENTE QUE LLEGASE A CAUSAR DAÑO A LOS BIENES '
              'O AL PERSONAL DE "EL PATRON".'),
            ('VIGESIMA CUARTA.', ' "EL TRABAJADOR" MANIFIESTA CONOCER '
              'PERFECTAMENTE EL ALCANCE Y PROCEDIMIENTOS DEL AVISO DE '
              'PRIVACIDAD QUE OPERA EN LA EMPRESA ESTEVEZ.JOR SERVICIOS S.A '
              'DE C.V PARA EFECTOS DE QUE SE VEA SALVAGUARDADA LA INFORMACION '
              'DE SUS DATOS PERSONALES.'),
            ('VIGESIMA QUINTA.', ' "LAS PARTES" ESTABLECEN QUE "EL PATRON" '
              'RESPETA PLENAMENTE LA DIGNIDAD HUMANA, PROPICIA EL TRABAJO '
              'DIGNO Y DECENTE EN TODAS LAS RELACIONES LABORALES, PROMUEVE LA '
              'IGUALDAD SUSTANTIVA ENTRE TRABAJADORES Y TRABAJADORAS, CONFORME '
              'A LA NOM NMX-R-025-SCFI-2015.'),
            ('VIGESIMA SEXTA.', ' "EL PATRON" MANIFIESTA QUE TRATANDOSE DE '
              'MUJERES Y MENORES SE ESTARA A LO ESTABLECIDO EN EL TITULO '
              'QUINTO DE "LA LEY".'),
            ('VIGESIMA SEPTIMA.', ' "EN CUMPLIMIENTO A LO SEÑALADO POR LOS '
              'ARTÍCULOS 501 Y 25 DE LA LEY, EN ESTE ACTO MANIFIESTO QUE '
              'SEÑALO COMO MI LEGITIMO BENEFICIARIO DE TODOS LOS BIENES Y '
              'PRESTACIONES LEGALES A: '
              '___________________________________________.'),
            ('VIGESIMA OCTAVA.', ' EN TEMAS DE SEGURIDAD E HIGIENE, EL '
              'TRABAJADOR SE OBLIGA A OBSERVAR TODAS LAS MEDIDAS Y NORMAS QUE '
              'APLICAN EN LA EMPRESA PARA PREVENIR ACCIDENTES, ENFERMEDADES Y '
              'LAS MEDIDAS DE SEGURIDAD E HIGIENE.'),
            ('VIGESIMA NOVENA.', ' EL PATRON LE HACE SABER AL TRABAJADOR QUE '
              'TIENE TOTAL LIBERTAD EN INTEGRAR LAS COMISIONES QUE SE INTEGREN '
              'EN LA EMPRESA CON MOTIVO DE LA CORRECTA APLICACIÓN Y '
              'SUPERVISION DE LAS CONDICIONES GENERALES DE TRABAJO.'),
            ('TRIGESIMA.', ' EL PATRON, RESPETA, SE OBLIGA Y GARANTIZA LA '
              'LIBERTAD SINDICAL QUE EN TERMINOS DE LEY TIENE EL TRABAJADOR.'),
            ('TRIGESIMA PRIMERA.', ' EL PATRON MANIFIESTA QUE TRATANDOSE DE '
              '"REPARTO DE UTILIDADES", SI AL TRABAJADOR LE CORRESPONDE ALGUNA '
              'SUMA POR ESTE CONCEPTO LA MISMA LE SERA ENTREGADA EN LOS '
              'TIEMPOS QUE SEÑALA LA LEY.'),
            ('TRIGESIMA SEGUNDA.', ' EL PATRON MANIFIESTA QUE EL TRABAJADOR '
              'GOZARA LOS BENEFICIOS DE LA SEGURIDAD SOCIAL A TRAVES DEL IMSS, '
              'ASI COMO EN EL INFONAVIT, SAR E INFONACOT.'),
            ('TRIGESIMA TERCERA.', ' EL PATRON MANIFIESTA QUE EN EL CENTRO DE '
              'TRABAJO SE HA IMPLEMENTADO UN PROTOCOLO PARA PREVENIR LA '
              'DISCRIMINACIÓN POR RAZONES DE GÉNERO Y ATENCIÓN DE CASOS DE '
              'VIOLENCIA Y ACOSO U HOSTIGAMIENTO SEXUAL.'),
            ('TRIGESIMA CUARTA.', ' EL PATRON MANIFIESTA QUE CONFORME A LA '
              'LEY OTORGARA A EL TRABAJADOR PERMISO DE PATERNIDAD DE CINCO '
              'DÍAS LABORABLES CON GOCE DE SUELDO.'),
            ('TRIGESIMA QUINTA.', ' LAS PARTES ESTABLECEN QUE TRATANDOSE DE '
              'TELETRABAJO EN CASO DE QUE ESTE SE DE, SE ESTARA A LO DISPUESTO '
              'POR EL TITULO CUARTO, CAPITULO XII BIS DE LA LEY.'),
            ('TRIGESIMA SEXTA.', ' EL PATRON MANIFIESTA TENER IMPLEMENTADAS '
              'LAS DISPOSICIONES PREVISTAS POR LA NOM-035-STPS-2018 Y LA '
              'NOM-036-1-STPS-2018.'),
            ('TRIGESIMA SEPTIMA.', ' EL PATRON MANIFIESTA A EL TRABAJADOR QUE '
              'TIENE IMPLEMENTADO EN EL CENTRO DE TRABAJO MEDIDAS Y PROTOCOLOS '
              'SANITARIOS PARA LA PREVENCIÓN DE ENFERMEDADES.'),
            ('TRIGESIMA OCTAVA.', ' AL TÉRMINO DE LA CAPACITACIÓN INICIAL, DE '
              'NO ACREDITAR COMPETENCIAS "EL TRABAJADOR", A JUICIO "DEL '
              'PATRÓN", SE DARÁ POR TERMINADA LA RELACIÓN DE TRABAJO SIN '
              'RESPONSABILIDAD PARA ESTE.'),
            ('TRIGESIMA NOVENA.', ' LAS PARTES CONVIENEN QUE TODO LO NO '
              'ESTIPULADO EXPRESAMENTE EN ESTE CONTRATO Y QUE SE ENCUENTRE '
              'CONTEMPLADO EN LA LEY, REGLAMENTO, SUS LEYES SUPLETORIAS, '
              'JURISPRUDENCIA Y LA COSTUMBRE LO TIENEN COMO PACTADO.'),
            ('CUADRAGÉSIMA.', ' "LAS PARTES" CONVIENEN QUE EN CASO DE '
              'CONTROVERSIA SE ESTARÁ A LO ESTABLECIDO EN EL PRESENTE Y EN LO '
              'QUE FUERE OMISO SE ENTENDERÁ A LO DISPUESTO POR LA LEY FEDERAL '
              'DEL TRABAJO BAJO LA JURISDICCIÓN DEL TRIBUNAL LABORAL DEL '
              'ESTADO DE MEXICO, COMO PRIMERA OPCIÓN.'),
        ]
        for num, txt in clausulas:
            b.p(doc, [(num, True), (txt, False)], align=J)

        b.p(doc, [
            (f'LEÍDO QUE FUE POR LAS PARTES EL DOCUMENTO E IMPUESTOS DEL '
             f'CONTENIDO, LAS PARTES LO RATIFICAN EN TODAS Y CADA UNA DE SUS '
             f'PARTES Y LO SUSCRIBEN POR DUPLICADO EL {date_str}, QUEDANDO UN '
             f'EJEMPLAR EN PODER DE CADA PARTE.', False),
        ], align=J)

        LINE = '_' * 48
        b.sig_table(doc,
            left_lines=[
                ('"EL PATRON"', True), ('', False), ('', False),
                (LINE, False),
                ('LIC. EDWIN GONZALEZ SORIA', True),
                ('APODERADO LEGAL', True), (emp_name, True),
            ],
            right_lines=[
                ('"EL TRABAJADOR"', True), ('', False), ('', False),
                (LINE, False), (emp_name, True),
            ],
        )

        # ---- Manifestación de No Adeudo ----
        doc.add_page_break()
        b.heading(doc, 'MANIFESTACIÓN DE NO ADEUDO', underline=True)

        b.p(doc, [
            ('QUIEN SUSCRIBE, C ', False), (emp_name, False),
            (' POR MEDIO DEL PRESENTE, MANIFIESTO BAJO PROTESTA DE DECIR '
             'VERDAD QUE A LA PRESENTE FECHA ', False),
            ('NO TENGO ADEUDOS PENDIENTES', True),
            (' CON EL INFONAVIT, INFONACOT, O DESCUENTOS POR PAGO DE PENSIÓN '
             'ALIMENTICIA, MOTIVO POR EL QUE RELEVO A LA EMPRESA ', False),
            (co, True),
            (' DE LA OBLIGACIÓN DE REALIZAR DESCUENTO ALGUNO DE MI SALARIO '
             'POR TALES CONCEPTOS.', False),
        ], align=J)

        b.p(doc, [
            ('PARA EL CASO DE QUE CUENTE CON ADEUDOS PENDIENTES CON EL '
             'INFONAVIT, INFONACOT, O DESCUENTOS POR PAGO DE PENSIÓN '
             'ALIMENTICIA, AUTORIZO A LA EMPRESA ', False),
            (co, True),
            (' PARA EFECTUAR LOS DESCUENTOS CORRESPONDIENTES EN TÉRMINOS DE '
             'LO DISPUESTO POR LA FRACCIÓN III, V Y VII DEL ARTÍCULO 110 DE '
             'LA LEY FEDERAL DEL TRABAJO.', False),
        ], align=J)

        b.p(doc, [
            ('LO ANTERIOR LO MANIFIESTO BAJO PROTESTA DE DECIR VERDAD, '
             'SABEDOR DE LAS SANCIONES QUE, EN CASO DE INFORMACIÓN FALSA Y '
             'ENGAÑOSA, ESTABLECEN PARA EL DELITO DE FRAUDE LOS ARTÍCULOS 386, '
             '387, 388 Y 388 BIS DEL CÓDIGO PENAL FEDERAL.', False),
        ], align=J)

        b.p(doc, [
            (f'QUEDA AUTORIZADO EL EMPLEADOR Y/O PATRÓN ', False),
            (co, True),
            (' PARA UTILIZAR LA PRESENTE DECLARACIÓN PARA TODOS LOS FINES '
             'LEGALES Y ADMINISTRATIVOS A QUE HAYA LUGAR.', False),
        ], align=J)

        b.p(doc, [(f'ESTADO DE MEXICO A {date_str}', True)], align=C,
            space_before=20)

        b.sig_table(doc,
            left_lines=[('', False)],
            right_lines=[(LINE, False), (f'C. {emp_name}', True)],
            space_pt=60,
        )

        # ---- Plan de Previsión Social ----
        doc.add_page_break()
        b.p(doc, [('C. ', True), ('REPRESENTANTE LEGAL DE LA EMPRESA', True)])
        b.p(doc, [('C. ', True), (co, True)])

        b.p(doc, [
            ('El suscripto de la presente es trabajador de la empresa '
             'ESTEVEZ.JOR SERVICIOS, S.A. DE C.V., toda vez que tiene '
             'celebrado contrato individual del trabajo por tiempo CERO DIAS '
             'con esa empresa, mismo contrato que a la fecha se encuentra '
             'vigente.', False),
        ], align=J)

        b.p(doc, [
            ('Manifiesto que se me ha dado a conocer el Plan de Previsión '
             'Social de la empresa ESTEVEZ.JOR SERVICIOS, S.A. DE C.V., y '
             'estoy enterado de las condiciones del mismo; de manera voluntaria '
             'solicito participar de los beneficios previstos en el Plan a '
             'partir de esta fecha.', False),
        ], align=J)

        b.p(doc, [('Sin más por el momento quedo de usted', False)], align=C)
        b.p(doc, [('ATENTAMENTE', True)], align=C, space_before=20)
        b.p(doc, [(emp_name, True)], align=C, space_before=30)

        doc.add_page_break()
        b.p(doc, [(date_str, True)], align=R)
        b.p(doc, [('C. ', True), (emp_name, True)], space_before=20)

        b.p(doc, [
            ('Por medio del presente escrito y de conformidad a la solicitud '
             'presentada por usted para participar de los beneficios previstos '
             'en el Plan de Previsión Social de la empresa ESTEVEZ.JOR '
             'SERVICIOS, S.A. DE C.V. y, toda vez que usted cumple con los '
             'requisitos de elegibilidad establecidos en el referido Plan, '
             'comunico a usted que ha sido aceptado para gozar de los '
             'beneficios de dicho Plan con el carácter de participante.', False),
        ], align=J)

        b.p(doc, [('ATENTAMENTE', True)], align=C, space_before=30)
        b.p(doc, [('PATRON, REPRESENTANTE LEGAL', True)], align=C)

        return b.download(self.env, 'hr.contract', self.id,
                          f'Contrato - {self.employee_id.name}.docx',
                          b.to_bytes(doc))

    # --- Addendum ---

    def action_download_addendum_docx(self):
        self.ensure_one()
        b = _DocxBuilder
        emp_name = self.employee_id.name.upper() if self.employee_id else 'N/A'
        today = date.today()
        mes = _DocxBuilder.MESES_ES[today.month].lower().capitalize()

        doc = Document()
        b.apply_default_typography(doc)
        for sec in doc.sections:
            sec.top_margin = Inches(1)
            sec.bottom_margin = Inches(1)
            sec.left_margin = Inches(1.2)
            sec.right_margin = Inches(1.2)

        b.heading(doc, 'ADDENDUM AL CONTRATO INDIVIDUAL DE TRABAJO')

        b.p(doc, [
            ('ADDENDUM AL CONTRATO INDIVIDUAL DE TRABAJO EL QUE CELEBRAN POR '
             'UNA PARTE LA SOCIEDAD DENOMINADA "ESTEVEZ.JOR SERVICIOS S.A. DE '
             'C.V..", EN ADELANTE EL "PATRÓN", REPRESENTADA EN ESTE ACTO POR '
             'EL C. EDWIN GONZALEZ SORIA Y POR OTRA PARTE EL C. ', False),
            (emp_name, True),
            (' A QUIEN EN LO SUCESIVO SE LE DENOMINARA EL "TRABAJADOR" '
             'PUDIENDO REFERIR A AMBOS COMO LAS "PARTES" AL TENOR DE LAS '
             'SIGUIENTES DECLARACIONES Y CLAUSULAS:', False),
        ], align=J)

        b.heading(doc, 'DECLARACIONES')
        b.p(doc, [('A.', True),
                  (' EL "PATRÓN" POR CONDUCTO DE SU REPRESENTANTE LEGAL SEÑALA '
                   'QUE:', False)])

        patron = [
            ('1.', ' QUE ES UNA SOCIEDAD ANÓNIMA DE CAPITAL VARIABLE DEBIDAMENTE '
              'CONSTITUIDA Y LEGALMENTE EXISTENTE DE ACUERDO CON LAS LEYES '
              'MEXICANAS, TAL Y COMO SE ACREDITA EN TÉRMINOS DE INSTRUMENTO '
              'NOTARIAL NÚMERO 46,404, DE FECHA 4 DE OCTUBRE DE 2011.'),
            ('2.', ' SU REPRESENTANTE LEGAL CUENTA CON PLENAS Y SUFICIENTES '
              'FACULTADES PARA CELEBRAR EL PRESENTE CONTRATO, TAL Y COMO LO '
              'ACREDITA CON EL INSTRUMENTO NOTARIAL NÚMERO 18,435, DE FECHA '
              '07 DE DICIEMBRE DE 2020.'),
            ('3.', ' SU REGISTRO FEDERAL DE CONTRIBUYENTES ES ESE111012E79, Y '
              'MANIFIESTA QUE SE ENCUENTRA AL CORRIENTE DEL CUMPLIMIENTO DE '
              'SUS OBLIGACIONES FISCALES.'),
            ('4.', ' QUE ES UNA SOCIEDAD ANÓNIMA DE CAPITAL VARIABLE DEBIDAMENTE '
              'CONSTITUIDA Y LEGALMENTE EXISTENTE DE ACUERDO CON LAS LEYES '
              'MEXICANAS (instrumento notarial número 46,404).'),
            ('5.', ' QUE SU REPRESENTADA CELEBRÓ CON LA PERSONA MORAL '
              '"MOVIMIENTO TECNOLOGICO, S. DE R.L. DE C.V." EN LO SUCESIVO EL '
              '"ADMINISTRADOR" UN CONTRATO DE PRESTACIÓN DE SERVICIOS RESPECTO '
              'DE UN PLAN INTEGRAL DE PREVISIÓN SOCIAL Y COMPENSACIONES '
              'REGISTRADO EN LA CONSAR.'),
        ]
        for num, txt in patron:
            b.p(doc, [(num, True), (txt, False)], align=J)

        b.p(doc, [('B.', True),
                  (' EL "TRABAJADOR" POR SU PROPIO DERECHO SEÑALA QUE:', False)],
            space_before=6)

        trabajador = [
            ('1.', ' TODAS Y CADA UNA DE LAS MANIFESTACIONES HECHAS EN EL '
              'CONTRATO INDIVIDUAL DE TRABAJO CELEBRADO CON EL "PATRÓN" SIGUEN '
              'SIENDO CIERTAS.'),
            ('2.', ' ESTÁ DE ACUERDO EN CELEBRAR EL PRESENTE INSTRUMENTO A '
              'EFECTO DE ANEXARLO AL CONTRATO INDIVIDUAL DE TRABAJO QUE TIENEN '
              'CELEBRADO CON EL "PATRÓN", A EFECTO DE QUE FORME PARTE INTEGRAL '
              'DE ÉSTE.'),
            ('3.', ' ASIMISMO, MANIFIESTA EL "TRABAJADOR" QUE, DE SER EL CASO, '
              'ESTÁ DE ACUERDO EN QUE SE DEN LOS AVISOS QUE SE REQUIERAN A LAS '
              'AUTORIDADES CORRESPONDIENTES PARA LOS EFECTOS LABORALES Y '
              'FISCALES CORRESPONDIENTES.'),
        ]
        for num, txt in trabajador:
            b.p(doc, [(num, True), (txt, False)], align=J)

        b.heading(doc, 'C L Á U S U L A S')

        clausulas = [
            ('PRIMERA.', ' - Que el "Trabajador" manifiesta que es su libre y '
              'entera voluntad que se reconfiguren sus percepciones que recibe '
              'por parte del "Patrón" para adherirse al Plan Integral de '
              'Previsión Social y Compensaciones, y que es de su pleno '
              'conocimiento que con dicho "Plan" se reconfiguren los conceptos '
              'que integran su percepción total a efecto de que con ello se '
              'aumente su percepción actual al obtener un 3% adicional como '
              'ahorro complementario para el retiro.'),
            ('SEGUNDA.', ' - En este acto el "Patrón" se da por enterado y '
              'debidamente notificado de la voluntad del "Trabajador" de '
              'adherirse al Plan Integral de Previsión Social y Compensaciones.'),
            ('TERCERA.', ' - Las "Partes", se obligan a guardar y mantener en '
              'secreto y con carácter de confidencial y no divulgar a terceras '
              'personas ningún tipo de "información confidencial" generada con '
              'motivo del presente instrumento.'),
            ('CUARTA.', ' - La duración del presente instrumento surte los '
              'mismos efectos del Contrato Individual de Trabajo que el '
              '"Trabajador" tiene celebrado con el "Patrón".'),
            ('QUINTA.', ' .- Finalmente, las "Partes" manifiestan y están de '
              'acuerdo en que el "Trabajador" conoce los términos y condiciones '
              'establecidos en los estatutos del "Plan Integral de Previsión '
              'Social y Compensaciones" y que firma y acepta conformidad la '
              'solicitud anexa al presente como Anexo "1".'),
            ('SEXTA.', f' - Que una vez requisitada la solicitud de mérito el '
              f'trabajador queda adherido al citado "Plan Integral de Previsión '
              f'Social y Compensaciones", y el presente instrumento es un '
              f'"Addendum" al contrato individual de trabajo. Leído que fue en '
              f'su totalidad el presente Addendum por las "Partes", lo firman '
              f'al margen y al calce en el Estado de México, el día {today.day} '
              f'del mes de {mes} del año {today.year} por no contener cláusula '
              f'alguna que atente contra el derecho, la moral o las buenas '
              f'costumbres.'),
        ]
        for num, txt in clausulas:
            b.p(doc, [(num, True), (txt, False)], align=J)

        b.p(doc, [('EL "TRABAJADOR"', True)], align=C, space_before=50)

        LINE = '_' * 48
        b.sig_table(doc,
            left_lines=[(LINE, False), ('POR SU PROPIO DERECHO', False)],
            right_lines=[(LINE, False), (emp_name, True)],
            space_pt=50,
        )

        return b.download(self.env, 'hr.contract', self.id,
                          f'Addendum - {self.employee_id.name}.docx',
                          b.to_bytes(doc))
